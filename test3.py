import os
import shutil
import threading
import traceback
import tkinter as tk
from tkinter import filedialog, messagebox, ttk

# --- MODIFIED IMPORTS to fix the TypeError ---
# Import the factory function with a unique name
from docx import Document as Document_factory
# Import the class types for type checking
from docx.document import Document
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.parts.image import ImagePart
from docx.shared import Inches
from docx.oxml.ns import qn


NSMAP = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
}

def iter_block_items(parent):
    # Use the imported Document class for the isinstance check
    if isinstance(parent, Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        # This case should ideally not be reached with correct usage
        raise ValueError(f"Parent must be a Document or _Cell object, but it is {type(parent)}")

    for child in parent_elm.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)

def copy_numbering(source_para, target_para):
    source_pPr = source_para._p.pPr
    if source_pPr is None:
        return
    numPr = source_pPr.find(qn("w:numPr"))
    if numPr is not None:
        target_pPr = target_para._p.get_or_add_pPr()
        existing_numPr = target_pPr.find(qn("w:numPr"))
        if existing_numPr is not None:
            target_pPr.remove(existing_numPr)
        target_pPr.append(numPr)

def copy_paragraph(source_para, target_container, temp_img_dir):
    if not source_para.text.strip() and not any(run._element.findall('.//w:drawing', namespaces=NSMAP) for run in source_para.runs):
         if hasattr(target_container, 'add_paragraph'):
            target_container.add_paragraph()
         return

    target_para = target_container.add_paragraph()
    target_para.style = source_para.style
    target_para.paragraph_format.alignment = source_para.paragraph_format.alignment
    copy_numbering(source_para, target_para)

    for run in source_para.runs:
        target_run = target_para.add_run(run.text)
        target_run.bold = run.bold
        target_run.italic = run.italic
        target_run.underline = run.underline
        if run.font:
            target_run.font.name = run.font.name
            target_run.font.size = run.font.size
            if run.font.color and run.font.color.rgb:
                target_run.font.color.rgb = run.font.color.rgb

        for drawing in run._element.findall(".//w:drawing", namespaces=NSMAP):
            rIds = drawing.findall('.//a:blip', namespaces=NSMAP)
            rId = rIds[0].get(qn("r:embed")) if rIds else None

            if rId:
                try:
                    related_part = source_para.part.related_parts[rId]
                    if isinstance(related_part, ImagePart):
                        image_bytes = related_part.blob
                        img_filename = f"{rId}.{related_part.default_ext}"
                        img_path = os.path.join(temp_img_dir, img_filename)

                        if not os.path.exists(img_path):
                            with open(img_path, "wb") as f:
                                f.write(image_bytes)

                        width = Inches(3)
                        try:
                            extent_elem = drawing.find('.//wp:extent', namespaces=NSMAP)
                            if extent_elem is not None:
                                cx = extent_elem.get('cx')
                                width = Inches(int(cx) / 914400)
                        except (AttributeError, ValueError):
                            pass

                        target_para.add_run().add_picture(img_path, width=width)
                    else:
                        placeholder_run = target_para.add_run("\n[Note: A complex graphical object (e.g., a chart or diagram) was here and could not be copied.]\n")
                        placeholder_run.font.italic = True
                except KeyError:
                    placeholder_run = target_para.add_run("\n[Note: A graphical object with a broken link was here and could not be copied.]\n")
                    placeholder_run.font.italic = True
            else:
                 placeholder_run = target_para.add_run("\n[Note: A shape or other drawing object was here and could not be copied.]\n")
                 placeholder_run.font.italic = True


def copy_table(source_table, target_container, temp_img_dir):
    # Use the imported Document class for the isinstance check
    if isinstance(target_container, Document):
         target_table = target_container.add_table(rows=0, cols=len(source_table.columns))
    elif isinstance(target_container, _Cell): # Handle nested tables
        placeholder = target_container.add_paragraph("[Nested Table Content Below]")
        try:
            placeholder.style = 'Caption'
        except KeyError: pass
        for nested_row in source_table.rows:
            row_text = "\t".join(cell.text.strip() for cell in nested_row.cells)
            target_container.add_paragraph(row_text)
        return
    else:
        return

    target_table.style = source_table.style
    target_table.autofit = source_table.autofit
    for i, col in enumerate(source_table.columns):
        if i < len(target_table.columns):
            target_table.columns[i].width = col.width

    for source_row in source_table.rows:
        target_row = target_table.add_row()
        for i, source_cell in enumerate(source_row.cells):
            target_cell = target_row.cells[i]
            target_cell._element.clear_content()
            for block in iter_block_items(source_cell):
                if isinstance(block, Paragraph):
                    copy_paragraph(block, target_cell, temp_img_dir)
                elif isinstance(block, Table):
                    copy_table(block, target_cell, temp_img_dir)


def split_docx_by_heading_with_images(docx_path, output_dir, heading_style="Heading 1", progress_callback=None):
    os.makedirs(output_dir, exist_ok=True)
    temp_img_dir = os.path.join(output_dir, "temp_images")
    os.makedirs(temp_img_dir, exist_ok=True)

    # Use the factory function to open a document
    doc = Document_factory(docx_path)
    sections = []
    blocks = list(iter_block_items(doc))

    if not blocks: return 0

    section_indices = []
    if not (isinstance(blocks[0], Paragraph) and blocks[0].style.name == heading_style):
        section_indices.append(0)

    for i, block in enumerate(blocks):
        if isinstance(block, Paragraph) and block.style.name == heading_style:
            section_indices.append(i)
    section_indices.append(len(blocks))

    for i in range(len(section_indices) - 1):
        start, end = section_indices[i], section_indices[i+1]
        if start == end: continue

        section_blocks = blocks[start:end]
        first_block = section_blocks[0]

        if isinstance(first_block, Paragraph) and first_block.style.name == heading_style:
            title = first_block.text.strip()
        else:
            title = "Introduction"
        if not title: title = f"Section_{i+1}"

        # Use the factory function to create a new document
        current_doc = Document_factory()
        if doc.sections:
            new_section = current_doc.sections[-1]
            ref_section = doc.sections[0]
            new_section.page_height, new_section.page_width = ref_section.page_height, ref_section.page_width
            new_section.left_margin, new_section.right_margin = ref_section.left_margin, ref_section.right_margin
            new_section.top_margin, new_section.bottom_margin = ref_section.top_margin, ref_section.bottom_margin

        for block in section_blocks:
            if isinstance(block, Paragraph):
                copy_paragraph(block, current_doc, temp_img_dir)
            elif isinstance(block, Table):
                copy_table(block, current_doc, temp_img_dir)
        sections.append((title, current_doc))

    total = len(sections)
    for i, (title, section_doc) in enumerate(sections):
        safe_title = "".join(c for c in title if c.isalnum() or c in " _-").rstrip()[:50]
        out_path = os.path.join(output_dir, f"{i+1:02d}_{safe_title}.docx")
        section_doc.save(out_path)
        if progress_callback:
            progress_callback(i + 1, total)

    if os.path.exists(temp_img_dir):
        shutil.rmtree(temp_img_dir)
    return total

def start_split_thread(filepath, output_dir):
    progress_bar['value'] = 0
    progress_bar['maximum'] = 100
    status_label.config(text="Splitting document...")
    button.config(state='disabled')

    def progress_callback(completed, total):
        if total > 0:
            percent = int(completed / total * 100)
            progress_bar['value'] = percent
            root.update_idletasks()

    def task():
        try:
            count = split_docx_by_heading_with_images(filepath, output_dir, progress_callback=progress_callback)
            if count > 0:
                messagebox.showinfo("Success", f"✅ Split complete.\n{count} sections saved in:\n{output_dir}")
            else:
                messagebox.showwarning("Warning", "No sections found based on 'Heading 1' style, or document is empty.")
        except Exception as e:
            print("An error occurred:")
            traceback.print_exc()
            messagebox.showerror("Error", f"❌ Error while splitting:\n{e}")
        finally:
            progress_bar['value'] = 0
            status_label.config(text="Ready")
            button.config(state='normal')

    threading.Thread(target=task, daemon=True).start()

def browse_and_split():
    filepath = filedialog.askopenfilename(
        title="Select Word File",
        filetypes=[("Word Documents", "*.docx")]
    )
    if not filepath: return

    filename = os.path.basename(filepath)
    name_only = os.path.splitext(filename)[0]
    output_dir = os.path.join(os.path.dirname(filepath), name_only + "_split")
    start_split_thread(filepath, output_dir)

# --- GUI Setup ---
root = tk.Tk()
root.title("DOCX Splitter by Heading")
root.geometry("400x180")
root.resizable(False, False)

style = ttk.Style(root)
style.theme_use('clam')

label = tk.Label(root, text="Split DOCX by 'Heading 1' sections", font=("Arial", 12, "bold"))
label.pack(pady=15)

button = ttk.Button(root, text="Select DOCX and Split", command=browse_and_split)
button.pack(pady=10, ipadx=10, ipady=5)

progress_bar = ttk.Progressbar(root, orient='horizontal', length=300, mode='determinate')
progress_bar.pack(pady=10)

status_label = tk.Label(root, text="Ready", font=("Arial", 10))
status_label.pack()

root.mainloop()