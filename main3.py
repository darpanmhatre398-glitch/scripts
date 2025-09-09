import os
import threading
import tkinter as tk
from tkinter import filedialog, messagebox
from tkinter import ttk
from pdf2docx import Converter
import fitz  # PyMuPDF
from docx import Document  # python-docx
import re


def convert_pdf_to_docx(pdf_path, docx_path):
    """Try converting PDF → DOCX with pdf2docx."""
    try:
        converter = Converter(pdf_path)
        converter.convert(docx_path)
        converter.close()
        return True
    except Exception as e:
        print(f"pdf2docx Conversion error: {e}")
        return False


def extract_text_to_docx(pdf_path, docx_path):
    """Fallback: extract structured text (paragraphs) from PDF into DOCX."""
    try:
        doc = fitz.open(pdf_path)
        word_doc = Document()

        def clean_text(s: str) -> str:
            # Remove NULLs and invalid control characters
            return re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F]', '', s)

        for page_num, page in enumerate(doc, start=1):
            word_doc.add_heading(f"Page {page_num}", level=2)

            # Extract text blocks (each block is like a paragraph)
            blocks = page.get_text("blocks")  # list of (x0, y0, x1, y1, text, block_no, block_type)
            if not blocks:
                word_doc.add_paragraph("[No extractable text]")
                continue

            # Sort by y-position, then x-position (top-to-bottom, left-to-right)
            blocks = sorted(blocks, key=lambda b: (b[1], b[0]))

            for b in blocks:
                text = clean_text(b[4].strip())
                if text:
                    word_doc.add_paragraph(text)

        word_doc.save(docx_path)
        return True
    except Exception as e:
        print(f"Text extraction error: {e}")
        return False


def run_conversion_thread(pdf_file, docx_file):
    progress_bar.start()
    convert_button.config(state=tk.DISABLED)

    success = False

    try:
        # Step 1: Try pdf2docx (full fidelity conversion)
        success = convert_pdf_to_docx(pdf_file, docx_file)

        # Step 2: If that fails, fallback to text extraction
        if not success:
            success = extract_text_to_docx(pdf_file, docx_file)

    except Exception as e:
        print(f"Processing error: {e}")
        success = False

    progress_bar.stop()
    convert_button.config(state=tk.NORMAL)

    if success and os.path.isfile(docx_file):
        messagebox.showinfo("Success", f"PDF successfully converted to:\n{docx_file}")
    else:
        messagebox.showerror("Error", "Conversion failed. See console for details.")


def browse_and_convert():
    pdf_file = filedialog.askopenfilename(
        title="Select PDF File",
        filetypes=[("PDF Files", "*.pdf")]
    )

    if not pdf_file:
        return

    docx_file = os.path.splitext(pdf_file)[0] + ".docx"

    if os.path.exists(docx_file):
        if not messagebox.askyesno("Overwrite?", f"{docx_file} already exists.\nDo you want to overwrite it?"):
            return

    threading.Thread(target=run_conversion_thread, args=(pdf_file, docx_file), daemon=True).start()


# GUI setup
root = tk.Tk()
root.title("PDF to DOCX Converter")
root.geometry("300x180")

label = tk.Label(root, text="Click the button to convert PDF to DOCX")
label.pack(pady=10)

convert_button = tk.Button(root, text="Select PDF and Convert", command=browse_and_convert)
convert_button.pack(pady=10)

progress_bar = ttk.Progressbar(root, mode='indeterminate')
progress_bar.pack(pady=10, fill=tk.X, padx=20)

root.mainloop()
