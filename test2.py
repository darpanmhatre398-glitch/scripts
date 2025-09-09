import os
import shutil
import threading
import traceback  # Import traceback for detailed error logging
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.shared import Inches
from docx.oxml.ns import qn

NSMAP = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
}

def iter_block_items(parent):
    if hasattr(parent, "element") and hasattr(parent.element, "body"):
        parent_elm = parent.element.body
    else:
        parent_elm = parent._element
    for child in parent_elm.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)

def copy_numbering(source_para, target_para):
    source_pPr = source_para._p.pPr
    if source_pPr is None:
        return
    numPr = source_pPr.find(qn("w:numPr"))
    if numPr is not None:
        target_pPr = target_para._p.get_or_add_pPr()
        existing_numPr = target_pPr.find(qn("w:numPr"))
        if existing_numPr is not None:
            target_pPr.remove(existing_numPr)
        target_pPr.append(numPr)

def copy_paragraph(source_para, target_container, temp_img_dir):
    # If the source paragraph is empty, add an empty one and return
    if not source_para.text.strip() and not source_para.runs:
         target_container.add_paragraph()
         return
         
    target_para = target_container.add_paragraph()
    target_para.style = source_para.style
    target_para.paragraph_format.alignment = source_para.paragraph_format.alignment
    copy_numbering(source_para, target_para)
    
    for run in source_para.runs:
        target_run = target_para.add_run(run.text)
        target_run.bold = run.bold
        target_run.italic = run.italic
        target_run.underline = run.underline
        if run.font:
            target_run.font.name = run.font.name
            target_run.font.size = run.font.size
            if run.font.color and run.font.color.rgb:
                target_run.font.color.rgb = run.font.color.rgb

        # --- Image Handling ---
        for drawing in run._element.findall(".//w:drawing", namespaces=NSMAP):
            for blip in drawing.findall(".//a:blip", namespaces=NSMAP):
                rId = blip.get(qn("r:embed"))
                if rId:
                    try:
                        image_part = source_para.part.related_parts[rId]
                        image_bytes = image_part.blob
                        # Create a unique filename for the image
                        img_filename = f"{rId}.{image_part.default_ext}"
                        img_path = os.path.join(temp_img_dir, img_filename)
                        if not os.path.exists(img_path):
                            with open(img_path, "wb") as f:
                                f.write(image_bytes)
                        
                        # Try to preserve original image size
                        cx = drawing.find('.//wp:extent', namespaces=drawing.nsmap).get('cx')
                        cy = drawing.find('.//wp:extent', namespaces=drawing.nsmap).get('cy')
                        width = Inches(int(cx) / 914400) if cx else Inches(3)

                        target_para.add_run().add_picture(img_path, width=width)
                    except (KeyError, AttributeError):
                        # Could not find or process the image, skip it
                        print(f"Warning: Could not process an image with rId {rId}.")
                        pass

def copy_table(source_table, target_container, temp_img_dir):
    # When target_container is a Document, use add_table
    if isinstance(target_container, type(Document())):
        target_table = target_container.add_table(rows=0, cols=len(source_table.columns))
    # When target_container is a _Cell, we can't add a table. We handle this below.
    else:
        # This is our workaround for nested tables
        placeholder = target_container.add_paragraph("[Nested Table Content Below]")
        try:
            placeholder.style = 'Caption'
        except KeyError:
            pass # Style 'Caption' may not exist
        for nested_row in source_table.rows:
            row_text = "\t".join(cell.text.strip() for cell in nested_row.cells)
            target_container.add_paragraph(row_text)
        return # Stop processing this table further

    target_table.style = source_table.style
    target_table.autofit = source_table.autofit
    # Copy column widths
    for i, col in enumerate(source_table.columns):
        if i < len(target_table.columns):
            target_table.columns[i].width = col.width

    for source_row in source_table.rows:
        target_row = target_table.add_row()
        for i, source_cell in enumerate(source_row.cells):
            target_cell = target_row.cells[i]
            target_cell._element.clear_content() # Clear the default paragraph
            for block in iter_block_items(source_cell):
                if isinstance(block, Paragraph):
                    copy_paragraph(block, target_cell, temp_img_dir)
                elif isinstance(block, Table):
                    # Recursive call now goes to our workaround logic
                    copy_table(block, target_cell, temp_img_dir)

def split_docx_by_heading_with_images(docx_path, output_dir, heading_style="Heading 1", progress_callback=None):
    if not os.path.exists(output_dir):
        os.makedirs(output_dir, exist_ok=True)
    temp_img_dir = os.path.join(output_dir, "temp_images")
    if not os.path.exists(temp_img_dir):
        os.makedirs(temp_img_dir, exist_ok=True)

    doc = Document(docx_path)
    sections = []
    
    # Collect all top-level blocks first
    blocks = list(iter_block_items(doc))
    if not blocks:
        return 0

    # Identify section breaks by heading style
    section_indices = []
    
    # Handle content before the first heading
    if not (isinstance(blocks[0], Paragraph) and blocks[0].style.name == heading_style):
        section_indices.append(0)

    for i, block in enumerate(blocks):
        if isinstance(block, Paragraph) and block.style.name == heading_style:
            section_indices.append(i)
    
    # Add an end marker for the last section
    section_indices.append(len(blocks))

    # Create sections based on indices
    for i in range(len(section_indices) - 1):
        start = section_indices[i]
        end = section_indices[i+1]
        
        section_blocks = blocks[start:end]
        first_block = section_blocks[0]
        
        # Determine the title for the file
        if isinstance(first_block, Paragraph) and first_block.style.name == heading_style:
            title = first_block.text.strip()
        else:
            title = "Introduction" # Default title for content before the first heading
        
        if not title:
            title = f"Section_{i+1}" # Fallback for empty titles

        current_doc = Document()
        # Copy document properties like page size
        for section in doc.sections:
             new_section = current_doc.sections[-1]
             new_section.page_height = section.page_height
             new_section.page_width = section.page_width
             new_section.left_margin = section.left_margin
             new_section.right_margin = section.right_margin
             new_section.top_margin = section.top_margin
             new_section.bottom_margin = section.bottom_margin
             break # Only need the first section's properties

        for block in section_blocks:
            if isinstance(block, Paragraph):
                copy_paragraph(block, current_doc, temp_img_dir)
            elif isinstance(block, Table):
                copy_table(block, current_doc, temp_img_dir)
        
        sections.append((title, current_doc))

    total = len(sections)
    for i, (title, section_doc) in enumerate(sections):
        safe_title = "".join(c for c in title if c.isalnum() or c in " _-").rstrip()[:50]
        out_path = os.path.join(output_dir, f"{i+1:02d}_{safe_title}.docx")
        section_doc.save(out_path)
        if progress_callback:
            progress_callback(i + 1, total)

    if os.path.exists(temp_img_dir):
        shutil.rmtree(temp_img_dir)
    return total

# GUI Functions with threading and progress bar

def start_split_thread(filepath, output_dir):
    progress_bar['value'] = 0
    progress_bar['maximum'] = 100
    status_label.config(text="Splitting document...")
    button.config(state='disabled')

    def progress_callback(completed, total):
        if total > 0:
            percent = int(completed / total * 100)
            progress_bar['value'] = percent
            root.update_idletasks()

    def task():
        try:
            count = split_docx_by_heading_with_images(filepath, output_dir, progress_callback=progress_callback)
            if count > 0:
                messagebox.showinfo("Success", f"✅ Split complete.\n{count} sections saved in:\n{output_dir}")
            else:
                messagebox.showwarning("Warning", "No sections found or document is empty.")
        except Exception as e:
            # Print detailed error to console and show a message box
            print("An error occurred:")
            traceback.print_exc()
            messagebox.showerror("Error", f"❌ Error while splitting:\n{e}")
        finally:
            progress_bar['value'] = 0
            status_label.config(text="Ready")
            button.config(state='normal')

    threading.Thread(target=task, daemon=True).start()

def browse_and_split():
    filepath = filedialog.askopenfilename(
        title="Select Word File",
        filetypes=[("Word Documents", "*.docx")]
    )
    if not filepath:
        return

    filename = os.path.basename(filepath)
    name_only = os.path.splitext(filename)[0]
    output_dir = os.path.join(os.path.dirname(filepath), name_only + "_split")

    start_split_thread(filepath, output_dir)

# GUI Setup
root = tk.Tk()
root.title("DOCX Splitter by Heading")
root.geometry("400x180")
root.resizable(False, False)

label = tk.Label(root, text="Split DOCX by 'Heading 1' sections", font=("Arial", 12))
label.pack(pady=15)

button = tk.Button(root, text="Select DOCX and Split", font=("Arial", 11), width=25, command=browse_and_split)
button.pack(pady=10)

progress_bar = ttk.Progressbar(root, orient='horizontal', length=300, mode='determinate')
progress_bar.pack(pady=10)

status_label = tk.Label(root, text="Ready", font=("Arial", 10))
status_label.pack()

root.mainloop()