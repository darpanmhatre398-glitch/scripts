import os
import re
import shutil
import tempfile
import threading
import queue
import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
from datetime import datetime

from docx import Document

# ==============================================================================
# 1. CORE LOGIC (Modified to accept and return sequence number)
# ==============================================================================

def generate_icn_code(dmc_code, kpc, xyz, sq, icv, issue, sec):
    parts = dmc_code.split("-")
    subsystem_index = -1
    for i in range(len(parts) - 1):
        if re.fullmatch(r"\d{2}", parts[i]) and re.fullmatch(r"\d{2}", parts[i + 1]):
            subsystem_index = i + 1
            break
    if subsystem_index == -1:
        return None
    up_to_subsystem = "-".join(parts[:subsystem_index + 1])
    return f"ICN-{up_to_subsystem}-{kpc}-{xyz}-{sq}-{icv}-{issue}-{sec}"

def update_doc_with_icn_labels(input_path, params, start_sq, pad_len, log_queue):
    dmc_code = params['dmc_code']
    log_queue.put(f"  • Using DMC: {dmc_code}")
    doc = Document(input_path)
    sq = start_sq  # Use the starting sequence number passed into the function
    paragraphs, i, generated_icns = doc.paragraphs[:], 0, []

    while i < len(paragraphs):
        para = paragraphs[i]
        has_image = any('graphic' in run._element.xml for run in para.runs)
        if has_image:
            caption_found = False
            if "figure" in para.text.lower():
                caption_found = True
            else:
                for j in range(1, 4):
                    if (i + j) < len(paragraphs):
                        caption_text = paragraphs[i + j].text.strip()
                        if caption_text:
                            if caption_text.lower().startswith("figure"):
                                caption_found = True
                            break
            if caption_found:
                icn = generate_icn_code(
                    dmc_code, params['kpc'], params['xyz'],
                    str(sq).zfill(pad_len), params['icv'], params['issue'], params['sec']
                )
                sq += 1  # Increment sequence number for each image
                if icn:
                    para._p.addnext(doc.add_paragraph(icn)._p)
                    generated_icns.append(icn)
                    paragraphs = doc.paragraphs[:]
                    i += 1
        i += 1
    
    log_queue.put(f"  • Found {len(generated_icns)} captioned images. Next SQ will be {sq}.")
    
    temp_output = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_output.name)
    # Return the temp file path AND the next sequence number to use
    return temp_output.name, sq

# ==============================================================================
# 2. TKINTER BATCH PROCESSING GUI
# ==============================================================================

class App:
    def __init__(self, root):
        self.root = root
        self.root.title("ICN Batch Processor (Continuous Sequence)")
        self.root.geometry("600x650")

        self.input_folder = tk.StringVar()
        self.output_folder = tk.StringVar()
        self.security_options = {
            "01-Unclassified": "01", "02-UK official Sensitive": "02",
            "03-RESTRICTED": "03", "02-INTERNAL": "02", "05-Confidential": "04"
        }

        main_frame = ttk.Frame(root, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)

        # Folder Selection
        file_frame = ttk.Labelframe(main_frame, text="1. Select Folders", padding="10")
        file_frame.pack(fill=tk.X, expand=True)
        ttk.Label(file_frame, text="Input Folder (contains .docx files):").grid(row=0, column=0, sticky="w", pady=(0,5))
        self.input_entry = ttk.Entry(file_frame, textvariable=self.input_folder, state="readonly")
        self.input_entry.grid(row=1, column=0, sticky="ew", padx=(0, 5))
        self.browse_input_btn = ttk.Button(file_frame, text="Browse...", command=lambda: self.select_folder(self.input_folder, "Select Input Folder"))
        self.browse_input_btn.grid(row=1, column=1)
        ttk.Label(file_frame, text="Output Folder (where results will be saved):").grid(row=2, column=0, sticky="w", pady=(5,5))
        self.output_entry = ttk.Entry(file_frame, textvariable=self.output_folder, state="readonly")
        self.output_entry.grid(row=3, column=0, sticky="ew", padx=(0, 5))
        self.browse_output_btn = ttk.Button(file_frame, text="Browse...", command=lambda: self.select_folder(self.output_folder, "Select Output Folder"))
        self.browse_output_btn.grid(row=3, column=1)
        file_frame.grid_columnconfigure(0, weight=1)

        # Parameters
        params_frame = ttk.Labelframe(main_frame, text="2. Set ICN Parameters", padding="10")
        params_frame.pack(fill=tk.X, expand=True, pady=10)
        self.entries = {}
        fields = {
            "RPC (KPC)": [str(i) for i in range(1, 10)], "XYZ (Origcage)": "1671Y",
            "Sequence Start": "00005", "Variant (ICV)": "A", "Issue": "001",
            "Security": list(self.security_options.keys())
        }
        for i, (label, value) in enumerate(fields.items()):
            ttk.Label(params_frame, text=f"{label}:").grid(row=i, column=0, sticky="w", padx=5, pady=5)
            if isinstance(value, list):
                self.entries[label] = ttk.Combobox(params_frame, values=value, state="readonly")
                self.entries[label].set(value[0])
            else:
                self.entries[label] = ttk.Entry(params_frame)
                self.entries[label].insert(0, value)
            self.entries[label].grid(row=i, column=1, sticky="ew", padx=5, pady=5)
        params_frame.grid_columnconfigure(1, weight=1)

        # Action Button
        self.run_btn = ttk.Button(main_frame, text="Start Batch Processing", command=self.start_processing)
        self.run_btn.pack(fill=tk.X, pady=10, ipady=5)

        # Log Window
        log_frame = ttk.Labelframe(main_frame, text="3. Status Log", padding="10")
        log_frame.pack(fill=tk.BOTH, expand=True)
        self.log_widget = scrolledtext.ScrolledText(log_frame, wrap=tk.WORD, height=10)
        self.log_widget.pack(fill=tk.BOTH, expand=True)
        self.log_widget.config(state="disabled")

        self.log_queue = queue.Queue()
        self.root.after(100, self.process_queue)

    def select_folder(self, string_var, title):
        path = filedialog.askdirectory(title=title)
        if path:
            string_var.set(path)

    def start_processing(self):
        input_dir = self.input_folder.get()
        output_dir = self.output_folder.get()
        if not input_dir or not output_dir:
            messagebox.showerror("Error", "Please select both an input and an output folder.")
            return

        self.run_btn.config(state="disabled")
        self.log_widget.config(state="normal")
        self.log_widget.delete('1.0', tk.END)
        self.log_widget.config(state="disabled")
        
        params = {
            'kpc': self.entries["RPC (KPC)"].get(), 'xyz': self.entries["XYZ (Origcage)"].get(),
            'sq_start': self.entries["Sequence Start"].get(), 'icv': self.entries["Variant (ICV)"].get(),
            'issue': self.entries["Issue"].get(), 'sec': self.security_options[self.entries["Security"].get()]
        }
        
        thread = threading.Thread(target=self.run_batch_thread, args=(input_dir, output_dir, params), daemon=True)
        thread.start()

    def run_batch_thread(self, input_dir, output_dir, params):
        try:
            timestamp = datetime.now().strftime("%Y-%m-%d_%H%M%S")
            final_output_folder = os.path.join(output_dir, f"output_{timestamp}")
            os.makedirs(final_output_folder, exist_ok=True)
            self.log_queue.put(f"✅ Created output folder: {final_output_folder}\n")

            docx_files = [f for f in os.listdir(input_dir) if f.lower().endswith('.docx') and not f.startswith('~')]
            if not docx_files:
                self.log_queue.put("⚠️ No .docx files found in the selected input folder.")
                self.log_queue.put("DONE")
                return
            
            # --- CONTINUOUS SEQUENCE LOGIC ---
            # Initialize sequence number and padding length from UI once
            current_sq = int(params['sq_start'])
            pad_len = len(params['sq_start'])
            
            for filename in docx_files:
                self.log_queue.put(f"📂 Processing: {filename}")
                input_path = os.path.join(input_dir, filename)
                
                current_params = params.copy()
                current_params['dmc_code'] = os.path.splitext(filename)[0]

                # Pass the current sequence number and get the next one back
                temp_path, next_sq = update_doc_with_icn_labels(input_path, current_params, current_sq, pad_len, self.log_queue)
                
                # Update the sequence number for the next file
                current_sq = next_sq
                
                final_path = os.path.join(final_output_folder, filename)
                shutil.copy(temp_path, final_path)
                os.remove(temp_path)
                self.log_queue.put(f"  -> Saved to output folder.\n")
            
            self.log_queue.put("🎉 Batch processing complete!")
        except Exception as e:
            self.log_queue.put(f"\n❌ AN UNEXPECTED ERROR OCCURRED: {e}")
        finally:
            self.log_queue.put("DONE")

    def process_queue(self):
        try:
            while True:
                msg = self.log_queue.get_nowait()
                if msg == "DONE":
                    self.run_btn.config(state="normal")
                    return
                self.log_widget.config(state="normal")
                self.log_widget.insert(tk.END, msg + "\n")
                self.log_widget.see(tk.END)
                self.log_widget.config(state="disabled")
        except queue.Empty:
            pass
        self.root.after(100, self.process_queue)

if __name__ == "__main__":
    root = tk.Tk()
    app = App(root)
    root.mainloop()